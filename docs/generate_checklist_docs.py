#!/usr/bin/env python3
"""
Generate the Excel workbook and Word document for the CardioX verification checklists.

The markdown files stay the single source of truth. Re-run this after editing them:

    python docs/generate_checklist_docs.py

Outputs (into docs/):
    CardioX_EXE_Test_Checklist.xlsx   fillable workbook, one row per check
    CardioX_EXE_Test_Checklist.docx   printable sign-off document
    CardioX_ReportScreen_Kotlin.docx  the ReportScreen.kt reference, as a Word doc
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from openpyxl import Workbook
from openpyxl.formatting.rule import CellIsRule
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation

import docx
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

DOCS = Path(__file__).resolve().parent
EXE_MD = DOCS / "EXE_TEST_CHECKLIST.md"
KT_MD = DOCS / "ReportScreen_kt.md"

# Layers that must be fully green before a build may ship (see the exit criteria).
BLOCKING = {"AUT", "BLD", "INS", "LNC", "LIC", "SEC"}

RESULTS = ["PASS", "FAIL", "N/A", "BLOCKED"]

# ── palette ──────────────────────────────────────────────────────────────────
NAVY = "1F3864"
BLUE = "2E5C8A"
LIGHT = "DCE6F1"
BAND = "F2F6FA"
GREEN = "C6EFCE"
RED = "FFC7CE"
GREY = "E7E6E6"
AMBER = "FFE699"
GREEN_TX = "006100"
RED_TX = "9C0006"

THIN = Side(style="thin", color="BFBFBF")
BOX = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)


# ── markdown helpers ─────────────────────────────────────────────────────────

_LINK = re.compile(r"\[([^\]]+)\]\([^)]+\)")
_BOLD = re.compile(r"\*\*(.+?)\*\*")
_CODE = re.compile(r"`([^`]+)`")
_ITALIC = re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)")


def md_to_text(s: str) -> str:
    """Flatten inline markdown to plain text."""
    s = _LINK.sub(r"\1", s)
    s = _BOLD.sub(r"\1", s)
    s = _CODE.sub(r"\1", s)
    s = _ITALIC.sub(r"\1", s)
    return re.sub(r"\s+", " ", s).strip()


def split_title(raw: str) -> tuple[str, str]:
    """
    Split a check into a short title and the remaining detail.

    A '**Lead-in.**' becomes the title when present; otherwise the first sentence is
    used, capped so the Excel column stays readable.
    """
    m = re.match(r"^\*\*(.+?)\*\*\s*(.*)$", raw, re.S)
    if m:
        return md_to_text(m.group(1)).rstrip("."), md_to_text(m.group(2))

    text = md_to_text(raw)
    m = re.match(r"^(.{0,110}?[.:])\s+(.*)$", text, re.S)
    if m and len(m.group(2)) > 12:
        return m.group(1).rstrip("."), m.group(2)
    if len(text) <= 130:
        return text.rstrip("."), ""
    cut = text.rfind(" ", 0, 130)
    return text[:cut].rstrip(" .,"), text[cut:].strip()


# ── parsing ──────────────────────────────────────────────────────────────────

H2 = re.compile(r"^##\s+(.*?)\s*$")
CHECK = re.compile(r"^-\s\[\s\]\s`([A-Z]{3,4})-(\d+)`\s+(.*)$")
ROW_CHECK = re.compile(r"^\|\s*-\s\[\s\]\s`([A-Z]{3,4})-(\d+)`\s*\|\s*(.*?)\s*\|\s*(.*?)\s*\|\s*$")
GROUP = re.compile(r"^\*\*(.+?)\*\*\s*$")


def clean_heading(h: str) -> str:
    h = re.sub(r"^\d+\.\s*", "", h)
    h = re.sub(r"\s*\(`[^`]+`(\s*/\s*`[^`]+`)?\)\s*$", "", h)
    h = re.sub(r"^Layer\s+\d+\s+[—-]\s*", "", h)
    return md_to_text(h)


def parse(path: Path) -> list[dict]:
    """Return one dict per check, in document order."""
    checks: list[dict] = []
    section = ""
    group = ""
    cur: dict | None = None

    for line in path.read_text(encoding="utf-8").splitlines():
        if m := H2.match(line):
            section, group, cur = clean_heading(m.group(1)), "", None
            continue

        if m := ROW_CHECK.match(line):
            prefix, num, desc, ref = m.groups()
            title, detail = split_title(desc)
            checks.append({
                "id": f"{prefix}-{num}", "prefix": prefix, "layer": section,
                "group": group, "title": title,
                "detail": (detail + (f"  Verified by: {md_to_text(ref)}." if ref else "")).strip(),
            })
            cur = None
            continue

        if m := CHECK.match(line):
            prefix, num, desc = m.groups()
            cur = {"id": f"{prefix}-{num}", "prefix": prefix, "layer": section,
                   "group": group, "raw": desc}
            checks.append(cur)
            continue

        if cur is not None and line.startswith(("      ", "\t")) and line.strip():
            cur["raw"] += " " + line.strip()
            continue

        if m := GROUP.match(line):
            group = md_to_text(m.group(1))
        if not line.strip():
            cur = None

    for c in checks:
        if "raw" in c:
            c["title"], c["detail"] = split_title(c.pop("raw"))
    return checks


# ── Excel ────────────────────────────────────────────────────────────────────

def style_header(ws, row: int, last_col: int, fill: str = NAVY) -> None:
    for c in range(1, last_col + 1):
        cell = ws.cell(row=row, column=c)
        cell.font = Font(bold=True, color="FFFFFF", size=10)
        cell.fill = PatternFill("solid", fgColor=fill)
        cell.alignment = Alignment(vertical="center", horizontal="center", wrap_text=True)
        cell.border = BOX
    ws.row_dimensions[row].height = 28


def title_block(ws, title: str, subtitle: str, span: int) -> int:
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=span)
    c = ws.cell(row=1, column=1, value=title)
    c.font = Font(bold=True, size=16, color=NAVY)
    c.alignment = Alignment(vertical="center")
    ws.row_dimensions[1].height = 26

    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=span)
    c = ws.cell(row=2, column=1, value=subtitle)
    c.font = Font(size=10, italic=True, color="595959")
    return 4


def checklist_sheet(wb: Workbook, name: str, title: str, subtitle: str,
                    checks: list[dict], group_col: bool) -> str:
    ws = wb.create_sheet(name)
    headers = ["ID", "Layer"] + (["Group"] if group_col else []) + \
              ["Blocking", "Check", "Detail / expected result", "Result",
               "Tester", "Date", "Defect #", "Notes"]
    hdr = title_block(ws, title, subtitle, len(headers))

    for i, h in enumerate(headers, start=1):
        ws.cell(row=hdr, column=i, value=h)
    style_header(ws, hdr, len(headers))

    r = hdr + 1
    first = r
    prev_layer = None
    for c in checks:
        if c["layer"] != prev_layer:
            ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=len(headers))
            cell = ws.cell(row=r, column=1, value=f"  {c['layer']}")
            cell.font = Font(bold=True, size=11, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor=BLUE)
            cell.alignment = Alignment(vertical="center")
            ws.row_dimensions[r].height = 20
            prev_layer = c["layer"]
            r += 1

        block_col = 4 if group_col else 3          # column holding "Blocking"
        vals = [c["id"], c["layer"]] + ([c["group"]] if group_col else []) + \
               ["YES" if c["prefix"] in BLOCKING else "", c["title"], c["detail"],
                "", "", "", "", ""]
        for i, v in enumerate(vals, start=1):
            cell = ws.cell(row=r, column=i, value=v)
            cell.border = BOX
            cell.alignment = Alignment(vertical="top", wrap_text=(i > block_col))
            if r % 2 == 0:
                cell.fill = PatternFill("solid", fgColor=BAND)
        ws.cell(row=r, column=1).font = Font(bold=True, size=10, name="Consolas")
        bl = ws.cell(row=r, column=block_col)
        if bl.value == "YES":
            bl.font = Font(bold=True, color=RED_TX)
            bl.alignment = Alignment(horizontal="center", vertical="top")
        r += 1

    last = r - 1
    res_col = headers.index("Result") + 1
    letter = get_column_letter(res_col)

    dv = DataValidation(type="list", formula1=f'"{",".join(RESULTS)}"', allow_blank=True)
    dv.prompt = "Select the outcome for this check"
    dv.promptTitle = "Result"
    ws.add_data_validation(dv)
    dv.add(f"{letter}{first}:{letter}{last}")

    rng = f"{letter}{first}:{letter}{last}"
    for value, fill, txt in (("PASS", GREEN, GREEN_TX), ("FAIL", RED, RED_TX),
                             ("N/A", GREY, "595959"), ("BLOCKED", AMBER, "7F6000")):
        ws.conditional_formatting.add(rng, CellIsRule(
            operator="equal", formula=[f'"{value}"'],
            fill=PatternFill("solid", bgColor=fill),
            font=Font(bold=True, color=txt)))

    widths = [11, 22] + ([18] if group_col else []) + [10, 46, 74, 11, 14, 12, 10, 30]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w

    ws.freeze_panes = ws.cell(row=first, column=1)
    ws.auto_filter.ref = f"A{hdr}:{get_column_letter(len(headers))}{last}"
    ws.sheet_view.zoomScale = 100
    return f"{name}!{letter}"


def build_workbook(exe_checks: list[dict], kt_checks: list[dict], out: Path) -> None:
    wb = Workbook()
    wb.remove(wb.active)

    # ── Read Me ──────────────────────────────────────────────────────────────
    ws = wb.create_sheet("Read Me")
    ws.column_dimensions["A"].width = 3
    ws.column_dimensions["B"].width = 26
    ws.column_dimensions["C"].width = 104

    ws["B2"] = "CardioX — EXE Release Verification"
    ws["B2"].font = Font(bold=True, size=18, color=NAVY)
    ws["B3"] = "Acceptance test workbook for the packaged executable and installer"
    ws["B3"].font = Font(size=11, italic=True, color="595959")

    rows = [
        ("How to use", "Work through the 'EXE Checklist' sheet top to bottom. Set a Result on every "
                       "row. The Summary sheet totals itself as you go."),
        ("Result values", "PASS · FAIL · N/A · BLOCKED — pick from the dropdown in the Result column. "
                          "Cells colour themselves."),
        ("Blocking layers", "AUT, BLD, INS, LNC, LIC and SEC must be 100% PASS. No exceptions — these "
                            "are ship-stopping."),
        ("Order matters", "Later sections assume earlier ones passed. Do not skip ahead."),
        ("Before you start", "Fill in the Environment sheet, and reset the machine: uninstall CardioX "
                             "and delete %LOCALAPPDATA%\\Deckmount\\CardioX."),
        ("Failures", "Log every FAIL on the Defect Log sheet and put its number in the Defect # column."),
        ("Android", "The 'Android (KRP)' sheet covers ReportScreen.kt — the Android report screen and "
                    "PDF export. It is tested separately from the desktop EXE."),
        ("Source of truth", "Generated from docs/EXE_TEST_CHECKLIST.md and docs/ReportScreen_kt.md. "
                            "Edit the markdown, then re-run docs/generate_checklist_docs.py."),
    ]
    r = 5
    for label, body in rows:
        ws.cell(row=r, column=2, value=label).font = Font(bold=True, size=10, color=NAVY)
        c = ws.cell(row=r, column=3, value=body)
        c.alignment = Alignment(wrap_text=True, vertical="top")
        c.font = Font(size=10)
        ws.row_dimensions[r].height = 30
        r += 1

    r += 1
    ws.cell(row=r, column=2, value="Sheets").font = Font(bold=True, size=12, color=NAVY)
    r += 1
    for nm, desc in (
        ("Environment", "Machine, build and tester details — fill in first"),
        ("EXE Checklist", f"{len(exe_checks)} checks across the desktop release"),
        ("Android (KRP)", f"{len(kt_checks)} checks for ReportScreen.kt"),
        ("Summary", "Live per-layer PASS/FAIL totals and the release verdict"),
        ("Defect Log", "One row per failure"),
        ("Sign-off", "Approval signatures"),
    ):
        ws.cell(row=r, column=2, value=nm).font = Font(bold=True, size=10, name="Consolas")
        ws.cell(row=r, column=3, value=desc).font = Font(size=10)
        r += 1

    # ── Environment ──────────────────────────────────────────────────────────
    ws = wb.create_sheet("Environment")
    ws.column_dimensions["A"].width = 3
    ws.column_dimensions["B"].width = 44
    ws.column_dimensions["C"].width = 60
    ws["B2"] = "Test environment"
    ws["B2"].font = Font(bold=True, size=16, color=NAVY)
    ws["B3"] = "A checklist result is meaningless without this."
    ws["B3"].font = Font(size=10, italic=True, color="595959")

    fields = [
        "Build version (src/version.py → APP_VERSION)", "Update channel",
        "Installer filename", "Installer SHA-256",
        "Build machine OS / Python version", "Test machine OS / build",
        "RAM / CPU / logical threads", "Environment matrix row (A minimum / B typical / C fresh)",
        "RhythmUltra serial under test", "Licence server URL baked into .env",
        "Tester name", "Date started", "Date completed",
    ]
    ws.cell(row=5, column=2, value="Field")
    ws.cell(row=5, column=3, value="Value")
    for c in (2, 3):
        cell = ws.cell(row=5, column=c)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor=NAVY)
        cell.border = BOX
        cell.alignment = Alignment(horizontal="center")
    for i, f in enumerate(fields):
        r = 6 + i
        a = ws.cell(row=r, column=2, value=f)
        a.font = Font(size=10)
        a.border = BOX
        a.alignment = Alignment(vertical="center", wrap_text=True)
        b = ws.cell(row=r, column=3)
        b.border = BOX
        b.fill = PatternFill("solid", fgColor=LIGHT)
        ws.row_dimensions[r].height = 20

    r = 6 + len(fields) + 2
    ws.cell(row=r, column=2, value="Environment matrix").font = Font(bold=True, size=12, color=NAVY)
    r += 1
    for row_name, spec in (
        ("A — Minimum", "<=8 GB RAM, i3 / <=4 logical threads — triggers is_low_spec_mode()"),
        ("B — Typical", "16 GB RAM, i5/i7, integrated GPU — the mainstream clinic machine"),
        ("C — Fresh Windows", "No Python, no MSVC redistributables — proves the bundle is complete"),
    ):
        ws.cell(row=r, column=2, value=row_name).font = Font(bold=True, size=10)
        ws.cell(row=r, column=3, value=spec).font = Font(size=10)
        ws.cell(row=r, column=3).alignment = Alignment(wrap_text=True)
        r += 1

    # ── the two checklists ───────────────────────────────────────────────────
    exe_ref = checklist_sheet(
        wb, "EXE Checklist", "CardioX EXE — release verification",
        "Set a Result on every row. Blocking layers must be 100% PASS.",
        exe_checks, group_col=False)
    checklist_sheet(
        wb, "Android (KRP)", "ReportScreen.kt — Android report screen",
        "Verification for the Android ECG report renderer and PDF export.",
        kt_checks, group_col=True)

    # ── Summary ──────────────────────────────────────────────────────────────
    ws = wb.create_sheet("Summary")
    ws.column_dimensions["A"].width = 3
    ws.column_dimensions["B"].width = 34
    for col in "CDEFGH":
        ws.column_dimensions[col].width = 13

    ws["B2"] = "Release summary"
    ws["B2"].font = Font(bold=True, size=16, color=NAVY)
    ws["B3"] = "Formulas are live — these totals update as the checklist is filled in."
    ws["B3"].font = Font(size=10, italic=True, color="595959")

    heads = ["Layer", "Blocking", "Checks", "PASS", "FAIL", "N/A", "Outstanding"]
    for i, h in enumerate(heads, start=2):
        ws.cell(row=5, column=i, value=h)
    style_header(ws, 5, 8)

    layers: list[tuple[str, str]] = []
    for c in exe_checks:
        if not layers or layers[-1][0] != c["layer"]:
            layers.append((c["layer"], c["prefix"]))

    sheet, rescol = exe_ref.split("!")
    quoted = f"'{sheet}'"
    r = 6
    for layer, prefix in layers:
        ws.cell(row=r, column=2, value=layer).font = Font(size=10)
        blocking = "YES" if prefix in BLOCKING else ""
        b = ws.cell(row=r, column=3, value=blocking)
        b.alignment = Alignment(horizontal="center")
        if blocking:
            b.font = Font(bold=True, color=RED_TX)
        crit = f'{quoted}!$B:$B,$B{r}'
        ws.cell(row=r, column=4, value=f'=COUNTIFS({crit})')
        for i, v in enumerate(("PASS", "FAIL", "N/A"), start=5):
            ws.cell(row=r, column=i,
                    value=f'=COUNTIFS({crit},{quoted}!${rescol}:${rescol},"{v}")')
        ws.cell(row=r, column=8, value=f"=D{r}-E{r}-F{r}-G{r}")
        for c in range(2, 9):
            ws.cell(row=r, column=c).border = BOX
            if c >= 4:
                ws.cell(row=r, column=c).alignment = Alignment(horizontal="center")
        r += 1

    tot = r
    ws.cell(row=tot, column=2, value="TOTAL")
    for c in range(4, 9):
        col = get_column_letter(c)
        ws.cell(row=tot, column=c, value=f"=SUM({col}6:{col}{tot - 1})")
    for c in range(2, 9):
        cell = ws.cell(row=tot, column=c)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor=NAVY)
        cell.border = BOX
        if c >= 4:
            cell.alignment = Alignment(horizontal="center")

    rng = f"E6:F{tot}"
    ws.conditional_formatting.add(rng, CellIsRule(
        operator="greaterThan", formula=["0"],
        fill=PatternFill("solid", bgColor=GREEN), font=Font(bold=True, color=GREEN_TX)))
    ws.conditional_formatting.add(f"F6:F{tot}", CellIsRule(
        operator="greaterThan", formula=["0"],
        fill=PatternFill("solid", bgColor=RED), font=Font(bold=True, color=RED_TX)))

    v = tot + 2
    ws.cell(row=v, column=2, value="Blocking layers clear?").font = Font(bold=True, size=11, color=NAVY)
    blocking_rows = [6 + i for i, (_, p) in enumerate(layers) if p in BLOCKING]
    fails = "+".join(f"F{i}" for i in blocking_rows)
    left = "+".join(f"H{i}" for i in blocking_rows)
    ws.cell(row=v, column=4,
            value=f'=IF(({fails})>0,"NO — blocking failures",'
                  f'IF(({left})>0,"NOT YET — checks outstanding","YES"))')
    ws.merge_cells(start_row=v, start_column=4, end_row=v, end_column=8)
    cell = ws.cell(row=v, column=4)
    cell.font = Font(bold=True, size=11)
    cell.alignment = Alignment(horizontal="center")
    cell.border = BOX
    ws.conditional_formatting.add(f"D{v}", CellIsRule(
        operator="equal", formula=['"YES"'],
        fill=PatternFill("solid", bgColor=GREEN), font=Font(bold=True, color=GREEN_TX)))
    ws.conditional_formatting.add(f"D{v}", CellIsRule(
        operator="notEqual", formula=['"YES"'],
        fill=PatternFill("solid", bgColor=AMBER), font=Font(bold=True, color="7F6000")))

    v += 2
    ws.cell(row=v, column=2, value="Release decision").font = Font(bold=True, size=12, color=NAVY)
    v += 1
    for label in ("APPROVED — all blocking layers green, remaining defects accepted",
                  "REJECTED — blocking failures recorded in the Defect Log"):
        ws.cell(row=v, column=2, value="☐")
        ws.cell(row=v, column=2).alignment = Alignment(horizontal="right")
        ws.cell(row=v, column=3, value=label).font = Font(size=10)
        v += 1

    # ── Defect Log ───────────────────────────────────────────────────────────
    ws = wb.create_sheet("Defect Log")
    heads = ["#", "Check ID", "Severity", "Summary", "Steps to reproduce",
             "Raised by", "Date", "Status"]
    ws.cell(row=1, column=1, value="Defect log").font = Font(bold=True, size=16, color=NAVY)
    ws.cell(row=2, column=1,
            value="Severity — Blocker: ship-stopping · Major: workaround exists · Minor: cosmetic")
    ws.cell(row=2, column=1).font = Font(size=10, italic=True, color="595959")
    for i, h in enumerate(heads, start=1):
        ws.cell(row=4, column=i, value=h)
    style_header(ws, 4, len(heads))

    dv_sev = DataValidation(type="list", formula1='"Blocker,Major,Minor"', allow_blank=True)
    dv_st = DataValidation(type="list", formula1='"Open,In progress,Fixed,Accepted,Closed"',
                           allow_blank=True)
    ws.add_data_validation(dv_sev)
    ws.add_data_validation(dv_st)
    for r in range(5, 45):
        ws.cell(row=r, column=1, value=r - 4).font = Font(size=10, color="808080")
        for c in range(1, len(heads) + 1):
            ws.cell(row=r, column=c).border = BOX
            ws.cell(row=r, column=c).alignment = Alignment(vertical="top", wrap_text=c in (4, 5))
    dv_sev.add(f"C5:C44")
    dv_st.add(f"H5:H44")
    ws.conditional_formatting.add("C5:C44", CellIsRule(
        operator="equal", formula=['"Blocker"'],
        fill=PatternFill("solid", bgColor=RED), font=Font(bold=True, color=RED_TX)))
    for i, w in enumerate([6, 12, 12, 44, 60, 16, 12, 14], start=1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = "A5"

    # ── Sign-off ─────────────────────────────────────────────────────────────
    ws = wb.create_sheet("Sign-off")
    ws.column_dimensions["A"].width = 3
    for col, w in zip("BCDE", (26, 30, 30, 18)):
        ws.column_dimensions[col].width = w
    ws["B2"] = "Release sign-off"
    ws["B2"].font = Font(bold=True, size=16, color=NAVY)
    ws["B3"] = "Sign only when the Summary sheet reports blocking layers clear."
    ws["B3"].font = Font(size=10, italic=True, color="595959")

    for i, h in enumerate(["Role", "Name", "Signature", "Date"], start=2):
        ws.cell(row=5, column=i, value=h)
    style_header(ws, 5, 5)
    for i, role in enumerate(("QA Engineer", "Engineering Lead", "Clinical Reviewer",
                              "Release Manager")):
        r = 6 + i
        ws.cell(row=r, column=2, value=role).font = Font(bold=True, size=10)
        for c in range(2, 6):
            ws.cell(row=r, column=c).border = BOX
        ws.row_dimensions[r].height = 34

    wb.save(out)


# ── Word ─────────────────────────────────────────────────────────────────────

def add_runs(p, raw: str) -> None:
    """Render inline markdown (**bold**, `code`) as Word runs."""
    for part in re.split(r"(\*\*.+?\*\*|`[^`]+`)", _LINK.sub(r"\1", raw)):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            p.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        else:
            p.add_run(part)


def shade(cell, hexcolor: str) -> None:
    tc = cell._tc.get_or_add_tcPr()
    el = docx.oxml.parse_xml(
        f'<w:shd {docx.oxml.ns.nsdecls("w")} w:fill="{hexcolor}"/>')
    tc.append(el)


COLS = (("ID", 0.72), ("Check", 2.85), ("Detail / expected result", 5.1),
        ("P", 0.33), ("F", 0.33))


def build_word_checklist(checks: list[dict], out: Path) -> None:
    doc = docx.Document()

    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Inches(0.6))

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.5)

    t = doc.add_heading("CardioX — EXE Release Verification Checklist", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = doc.add_paragraph()
    r = p.add_run("Acceptance test record for the packaged executable and installer")
    r.italic = True
    r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    doc.add_paragraph()
    meta = doc.add_table(rows=4, cols=4)
    meta.style = "Table Grid"
    for i, (a, b) in enumerate((("Build version", ""), ("Installer SHA-256", ""),
                                ("Test machine / matrix row", ""), ("Tester", ""))):
        meta.cell(i, 0).text = a
        meta.cell(i, 0).paragraphs[0].runs[0].bold = True
        shade(meta.cell(i, 0), LIGHT)
        meta.cell(i, 2).text = ("Date" if i == 0 else "Channel" if i == 1
                                else "RhythmUltra serial" if i == 2 else "Signature")
        meta.cell(i, 2).paragraphs[0].runs[0].bold = True
        shade(meta.cell(i, 2), LIGHT)

    doc.add_paragraph()
    h = doc.add_paragraph()
    add_runs(h, "**How to use.** Work top to bottom — later sections assume earlier ones passed. "
                "Tick one box per row. Layers **AUT, BLD, INS, LNC, LIC and SEC are blocking**: they "
                "must be 100% PASS before the build ships. Log every failure in the defect table at "
                "the end and reference its number.")

    prev = None
    for c in checks:
        if c["layer"] != prev:
            if prev is not None:
                doc.add_paragraph()
            head = doc.add_heading(c["layer"], level=1)
            if c["prefix"] in BLOCKING:
                run = head.add_run("   [BLOCKING]")
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x9C, 0x00, 0x06)

            tbl = doc.add_table(rows=1, cols=5)
            tbl.style = "Table Grid"
            tbl.autofit = False
            for i, (txt, w) in enumerate(COLS):
                cell = tbl.cell(0, i)
                cell.text = txt
                cell.width = Inches(w)
                para = cell.paragraphs[0]
                para.runs[0].bold = True
                para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                para.runs[0].font.size = Pt(9)
                shade(cell, NAVY)
            prev = c["layer"]

        row = tbl.add_row()
        for i, (_, w) in enumerate(COLS):
            row.cells[i].width = Inches(w)
        idc = row.cells[0].paragraphs[0].add_run(c["id"])
        idc.bold = True
        idc.font.name = "Consolas"
        idc.font.size = Pt(8.5)

        add_runs(row.cells[1].paragraphs[0], c["title"])
        for run in row.cells[1].paragraphs[0].runs:
            run.font.size = Pt(9)

        add_runs(row.cells[2].paragraphs[0], c["detail"])
        for run in row.cells[2].paragraphs[0].runs:
            run.font.size = Pt(8.5)

        for i in (3, 4):
            para = row.cells[i].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run("☐").font.size = Pt(12)

    doc.add_page_break()
    doc.add_heading("Defect log", level=1)
    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = "Table Grid"
    for i, txt in enumerate(("#", "Check ID", "Severity", "Summary",
                             "Steps to reproduce", "Status")):
        cell = tbl.cell(0, i)
        cell.text = txt
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(cell, NAVY)
    for n in range(1, 13):
        tbl.add_row().cells[0].text = str(n)

    doc.add_paragraph()
    doc.add_heading("Release decision", level=1)
    for label in ("APPROVED — all blocking layers green, remaining defects accepted below",
                  "REJECTED — blocking failures recorded in the defect log"):
        p = doc.add_paragraph()
        p.add_run("☐  ").font.size = Pt(13)
        p.add_run(label).bold = True

    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    for i, txt in enumerate(("Role", "Name", "Signature", "Date")):
        cell = tbl.cell(0, i)
        cell.text = txt
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(cell, NAVY)
    for role in ("QA Engineer", "Engineering Lead", "Clinical Reviewer", "Release Manager"):
        cells = tbl.add_row().cells
        cells[0].text = role
        cells[0].paragraphs[0].runs[0].bold = True

    doc.save(out)


def build_word_reference(md: Path, out: Path) -> None:
    """Convert the ReportScreen.kt reference markdown into a Word document."""
    doc = docx.Document()
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(doc.sections[0], m, Inches(0.8))
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    lines = md.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code: list[str] = []
    table: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table
        if len(table) < 2:
            table = []
            return
        header, body = table[0], [r for r in table[1:]
                                  if not all(set(c) <= set("-: ") for c in r)]
        tbl = doc.add_table(rows=1, cols=len(header))
        tbl.style = "Table Grid"
        for c, txt in enumerate(header):
            cell = tbl.cell(0, c)
            cell.text = ""
            add_runs(cell.paragraphs[0], txt)
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)
            shade(cell, NAVY)
        for row in body:
            cells = tbl.add_row().cells
            for c, txt in enumerate(row[:len(header)]):
                cells[c].text = ""
                add_runs(cells[c].paragraphs[0], txt)
                for run in cells[c].paragraphs[0].runs:
                    run.font.size = Pt(9)
        doc.add_paragraph()
        table = []

    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                r = p.add_run("\n".join(code))
                r.font.name = "Consolas"
                r.font.size = Pt(8.5)
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(10)
                code, in_code = [], False
            else:
                flush_table()
                in_code = True
            i += 1
            continue
        if in_code:
            code.append(line)
            i += 1
            continue

        if line.startswith("|"):
            table.append([c.strip() for c in line.strip().strip("|").split("|")])
            i += 1
            continue
        flush_table()

        if line.startswith("# "):
            doc.add_heading(md_to_text(line[2:]), level=0)
        elif line.startswith("## "):
            doc.add_heading(md_to_text(line[3:]), level=1)
        elif line.startswith("### "):
            doc.add_heading(md_to_text(line[4:]), level=2)
        elif line.startswith("---"):
            pass
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_runs(p, line[2:])
            for run in p.runs:
                run.italic = True
        elif re.match(r"^-\s\[\s\]\s", line):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.add_run("☐  ")
            add_runs(p, line[6:])
            for run in p.runs:
                run.font.size = Pt(9.5)
        elif line.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, line[2:])
        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\.\s", "", line))
        elif line.strip():
            add_runs(doc.add_paragraph(), line)
        i += 1

    flush_table()
    doc.save(out)


def main() -> int:
    exe_checks = parse(EXE_MD)
    kt_checks = parse(KT_MD)
    print(f"parsed {len(exe_checks)} EXE checks, {len(kt_checks)} Android checks")

    xlsx = DOCS / "CardioX_EXE_Test_Checklist.xlsx"
    build_workbook(exe_checks, kt_checks, xlsx)
    print(f"wrote {xlsx.name}")

    docx_out = DOCS / "CardioX_EXE_Test_Checklist.docx"
    build_word_checklist(exe_checks, docx_out)
    print(f"wrote {docx_out.name}")

    ref_out = DOCS / "CardioX_ReportScreen_Kotlin.docx"
    build_word_reference(KT_MD, ref_out)
    print(f"wrote {ref_out.name}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
